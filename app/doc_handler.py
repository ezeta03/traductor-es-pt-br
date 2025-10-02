from docx import Document
import fitz  # PyMuPDF

def extract_text_from_docx(path: str):
    doc = Document(path)
    content = []
    for para in doc.paragraphs:
        runs = []
        for run in para.runs:
            runs.append(run.text)
        content.append({"runs": runs})
    return content

def translate_docx(input_path: str, output_path: str, translator_func):
    doc = Document(input_path)

    def translate_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                text = run.text.strip()
                if text:
                    run.text = translator_func(text)

    def translate_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    # Traducir párrafos de la celda
                    translate_paragraphs(cell.paragraphs)
                    # Traducir tablas anidadas
                    if cell.tables:
                        translate_tables(cell.tables)

    # Traducir todo el documento
    translate_paragraphs(doc.paragraphs)
    translate_tables(doc.tables)

    doc.save(output_path)

def extract_text_from_pdf(path: str) -> str:
    doc = fitz.open(path)
    texts = []
    for page in doc:
        texts.append(page.get_text("text"))
    return "\n".join(texts)
