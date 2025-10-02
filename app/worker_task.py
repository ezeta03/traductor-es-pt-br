from celery import Celery
from docx import Document
import os
from app.translator_service import translate_text

celery = Celery(
    "worker_task",
    broker="redis://redis:6379/0",
    backend="redis://redis:6379/0"
)

# translator = Translator()

@celery.task
def translate_text_task(text: str, source: str, target: str) -> str:
    if not text.strip():
        return ""
    return translate_text(text, source, target)   # 👈 llamamos al servicio dinámico

@celery.task
def translate_file_task(input_path: str, output_path: str, source: str, target: str):
    if not os.path.exists(input_path):
        return {"status": "error", "error": "Archivo no existe"}

    doc = Document(input_path)
    translated_doc = Document()

    for para in doc.paragraphs:
        text = para.text
        if text.strip():
            translated_text = translate_text_task.delay(text, source, target).get(timeout=10)
        else:
            translated_text = ""

        p_new = translated_doc.add_paragraph(translated_text)

        # Copiar estilos
        if para.runs:
            for run_src, run_dest in zip(para.runs, p_new.runs):
                run_dest.bold = run_src.bold
                run_dest.italic = run_src.italic
                run_dest.underline = run_src.underline
                if run_src.font.name:
                    run_dest.font.name = run_src.font.name
                if run_src.font.size:
                    run_dest.font.size = run_src.font.size

    translated_doc.save(output_path)
    return {"status": "ok"}
