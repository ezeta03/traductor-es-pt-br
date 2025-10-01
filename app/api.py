from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
import shutil, os
from uuid import uuid4
from app.worker_task import translate_file_task

app = FastAPI()
UPLOAD_DIR = os.environ.get("UPLOAD_DIR", "/data/uploads")
OUT_DIR = os.environ.get("OUT_DIR", "/data/out")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUT_DIR, exist_ok=True)

@app.post("/translate")
async def translate_endpoint(file: UploadFile = File(...)):
    uid = str(uuid4())
    infile = os.path.join(UPLOAD_DIR, f"{uid}_{file.filename}")
    outname = os.path.join(OUT_DIR, f"{uid}_translated.docx")

    # Guardar archivo subido
    with open(infile, "wb") as f:
        shutil.copyfileobj(file.file, f)
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
from docx import Document
import os
from pathlib import Path

app = FastAPI()

UPLOAD_DIR = Path("data/in")
OUTPUT_DIR = Path("data/out")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

@app.post("/translate-docx/")
async def translate_docx(file: UploadFile = File(...)):
    # Guardar archivo subido
    input_path = UPLOAD_DIR / file.filename
    with open(input_path, "wb") as f:
        f.write(await file.read())

    # Abrir el DOCX
    doc = Document(input_path)

    # Crear documento de salida
    output_path = OUTPUT_DIR / f"translated_{file.filename}"
    translated_doc = Document()

    # Traducir párrafos (ejemplo con texto fijo)
    for para in doc.paragraphs:
        text = para.text
        if text.strip():
            # Aquí deberías llamar a la API de traducción real (ej: DeepL, Hugging Face, etc.)
            translated_text = text + " (traducido al portugués)"
        else:
            translated_text = ""

        p = translated_doc.add_paragraph()
        run = p.add_run(translated_text)
        run.bold = para.runs[0].bold if para.runs else False
        run.italic = para.runs[0].italic if para.runs else False

    translated_doc.save(output_path)

    # Retornar el archivo traducido
    return FileResponse(output_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=output_path.name)

    # Ejecutar traducción (sin Celery para respuesta directa)
    result = translate_file_task(infile, outname)

    if result.get("status") == "ok":
        return FileResponse(
            path=outname,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{os.path.splitext(file.filename)[0]}_translated.docx"
        )
    else:
        return {"status": "error", "error": result.get("error")}
